# exams/views.py
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate, login, logout
from django.contrib import messages
from django.utils import timezone
from django.urls import reverse
from django.http import HttpResponse, JsonResponse
from datetime import date, datetime,timedelta
from django.core.serializers.json import DjangoJSONEncoder
import json
import os
import io
from docx import Document
from django.conf import settings
from docx.shared import Inches, Pt,Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from django.urls import reverse
from django.utils.http import urlencode
from django.db import IntegrityError, transaction

from .models import ExamDuration, ExamBooking, AttendanceSheet,Subject,AdminResult, SEMESTER_CHOICES,CustomUser,ExamSchedule,RoomAssignment
from .forms import (
    StudentRegistrationForm,
    FacultyRegisterForm,
      FacultyLoginForm,
    ExamDurationForm,
    StudentLoginForm,
    AdminLoginForm,
    SubjectForm,
    AdminResultForm
)
from django.contrib.auth.hashers import make_password
from django.forms.models import model_to_dict


# -----------------------
# HOME
# -----------------------
def home(request):
    return render(request, "home.html")


# -----------------------
# ADMIN LOGIN
# -----------------------
def admin_login(request):
    if request.method == "POST":
        form = AdminLoginForm(request.POST)
        if form.is_valid():
            username = form.cleaned_data["username"]
            password = form.cleaned_data["password"]

            user = authenticate(request, username=username, password=password)
            if user and user.role == "admin":
                login(request, user)
                return redirect("admin_dashboard")
            else:
                messages.error(request, "Invalid credentials or not an admin.")
    else:
        form = AdminLoginForm()

    return render(request, "admin_login.html", {"form": form})


def admin_logout(request):
    logout(request)
    return redirect("admin_login")

# -----------------------
# FACULTY REGISTER
# -----------------------
def faculty_register(request):
    if request.method == "POST":
        form = FacultyRegisterForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "Registration successful! Wait for admin approval.")
            return redirect("faculty_login")
    else:
        form = FacultyRegisterForm()
    return render(request, "faculty_register.html", {"form": form})


# -----------------------
# FACULTY LOGIN
# -----------------------
def faculty_login(request):
    if request.method == "POST":
        form = FacultyLoginForm(request.POST)
        if form.is_valid():
            username = form.cleaned_data["username"]
            password = form.cleaned_data["password"]
            user = authenticate(username=username, password=password)

            if user is not None and user.role == "faculty":
                if user.is_approved:
                    login(request, user)
                    return redirect("faculty_dashboard")
                else:
                    messages.error(request, "Your account is not approved yet.")
            else:
                messages.error(request, "Invalid login details.")
    else:
        form = FacultyLoginForm()
    return render(request, "faculty_login.html", {"form": form})

# -----------------------
# STUDENT REGISTER
# -----------------------
def student_register(request):
    if request.method == "POST":
        form = StudentRegistrationForm(request.POST)
        if form.is_valid():
            # form.save() already handles password hashing
            form.save()  
            messages.success(request, "Registration successful! Please login.")
            return redirect("student_login")
        else:
            # Show form errors in template
            messages.error(request, "Please correct the errors below.")
    else:
        form = StudentRegistrationForm()

    return render(request, "student_register.html", {"form": form})
   

# -----------------------
# STUDENT LOGIN
# -----------------------
def student_login(request):
    if request.method == "POST":
        form = StudentLoginForm(request.POST)
        if form.is_valid():
            username = form.cleaned_data["username"]
            password = form.cleaned_data["password"]

            user = authenticate(request, username=username, password=password)
            if user and user.role == "student":
                login(request, user)
                return redirect("student_dashboard")
            else:
                messages.error(request, "Invalid credentials or not a student account.")
    else:
        form = StudentLoginForm()

    return render(request, "student_login.html", {"form": form})


@login_required
def student_logout(request):
    logout(request)
    return redirect("student_login")


# -----------------------
# STUDENT DASHBOARD
# -----------------------
@login_required(login_url="/student-login/")
def student_dashboard(request):
    if request.method == "POST":
        subject_id = request.POST.get("subject")
        date = request.POST.get("date")

        # Check if date is within any exam schedule
        schedule = ExamDuration.objects.filter(
            start_date__lte=date,
            end_date__gte=date
        ).first()

        if not schedule:
            messages.error(request, "Selected date is not within any exam schedule ❌")
            return redirect("student_dashboard")

        # Check if student has already booked this subject on this date
        if ExamBooking.objects.filter(student=request.user, subject_id=subject_id, date=date).exists():
            messages.error(request, "You have already booked this subject on this date ❌")
            return redirect("student_dashboard")

        # Check how many times this student has booked this subject in total
        total_booked_subject = ExamBooking.objects.filter(student=request.user, subject_id=subject_id).count()
        if total_booked_subject >= 2:
            messages.error(request, "This subject can be booked a maximum of 2 times ❌")
            return redirect("student_dashboard")

        # If all checks pass, create booking
        ExamBooking.objects.create(
            student=request.user,
            subject_id=subject_id,
            date=date,
            schedule=schedule
        )

        messages.success(request, "Exam booked successfully ✅")
        return redirect("student_dashboard")

    # GET request — show dashboard
    schedules = ExamDuration.objects.all()
    schedules_json = json.dumps(list(schedules.values("start_date", "end_date")), cls=DjangoJSONEncoder)

    # Only fetch subjects for student's semester
    subjects = Subject.objects.filter(semester=request.user.semester).order_by("code")

    return render(request, "student_dashboard.html", {
        "schedules_json": schedules_json,
        "student": request.user,
        "subjects_by_sem": json.dumps({request.user.semester: [model_to_dict(s) for s in subjects]}, cls=DjangoJSONEncoder)
    })


# -----------------------
# View Result
# -----------------------

@login_required
def view_results(request):
    student = request.user

    # Get results where student_id matches the logged-in student's ID
    results = AdminResult.objects.filter(student_id=student.student_id)

    context = {
        "student": student,
        "results": results,
    }

    return render(request, "student_results.html", context)

# -----------------------
# API - Exam Dates for Calendar
# -----------------------
@login_required
def get_exam_dates(request):
    duration = ExamDuration.objects.last()
    events = []

    if duration:
        current = duration.start_date
        while current <= duration.end_date:
            events.append({
                "title": "Available",
                "start": str(current),
                "allDay": True
            })
            current += timedelta(days=1)

    return JsonResponse(events, safe=False)


# -----------------------
# ADMIN DASHBOARD
# -----------------------
@login_required(login_url='/admin-login/')
def admin_dashboard(request):
    if request.user.role != "admin":
        messages.error(request, "Access Denied! Only Admin can access this page.")
        return redirect("student_dashboard")

    #  Total students who registered for any subject
    total_students = ExamBooking.objects.values('student').distinct().count()

    #  Total approved faculty
    total_faculty = CustomUser.objects.filter(role='faculty', is_approved=True).count()

    #  Total registrations (total bookings)
    total_registrations = ExamBooking.objects.count()

    #  Pending results (bookings with no result yet)
    # Approach: count bookings where no AdminResult exists for that student/subject/date
    from django.db.models import Q, F, Count, Subquery, OuterRef

    # Subquery to find results for each booking
    results_subquery = AdminResult.objects.filter(
        student_id=OuterRef('student__student_id'),
        subject_code=OuterRef('subject__code'),
        r_date=OuterRef('date')
    )

    pending_results = ExamBooking.objects.annotate(
        has_result=Subquery(results_subquery.values('id')[:1])
    ).filter(has_result__isnull=True).count()

    # Fetch all bookings for table display
    bookings = ExamBooking.objects.all().select_related('student', 'subject').order_by('date')

    return render(request, "admin_dashboard.html", {
        "total_students": total_students,
        "total_faculty": total_faculty,
        "total_registrations": total_registrations,
        "pending_results": pending_results,
        "bookings": bookings,
    })


# -----------------------
# SET SCHEDULE
# -----------------------
@login_required(login_url='/admin-login/')

def set_schedule(request):
    if request.method == "POST":
        form = ExamDurationForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "Exam schedule added successfully ✅")
            return redirect("set_schedule")
        else:
            messages.error(request, "Please correct the errors below ❌")
    else:
        form = ExamDurationForm()

    schedules = ExamDuration.objects.all().order_by("-start_date")
    return render(request, "set_schedule.html", {
        "form": form,
        "schedules": schedules
    })

def delete_schedule(request, schedule_id):
    schedule = get_object_or_404(ExamDuration, id=schedule_id)

    if request.method == "POST":
        try:
            with transaction.atomic():
                # 🔥 Delete ALL related RoomAssignments safely
                RoomAssignment.objects.filter(schedule=schedule).delete()

                # 🔥 Delete schedule (will delete other related objects via CASCADE)
                schedule.delete()

            messages.success(request, "✅ Schedule and all related records deleted successfully!")
        except Exception as e:
            messages.error(request, f"❌ Error deleting schedule: {str(e)}")

        return redirect("set_schedule")

    messages.warning(request, "⚠️ Invalid request method. Please use the delete button.")
    return redirect("set_schedule")
# -----------------------
# Attendance sheet
# -----------------------

# -------------------------
# Attendance Selector
# -------------------------
@login_required(login_url='/admin-login/')
def attendance_selector(request):
    today_date = date.today()  

    schedules = ExamDuration.objects.all()

    # Collect valid dates from all schedules
    valid_dates = []
    for schedule in schedules:
        for d in schedule.get_date_range():
            # Ensure comparison is date vs date
            if isinstance(d, datetime):
                d = d.date()
            if d >= today_date:  # skip past dates
                valid_dates.append(d)

    # Remove duplicates & sort
    valid_dates = sorted(set(valid_dates))

    # 🔹 Fetch subjects dynamically
    subjects = Subject.objects.all().order_by("semester", "code")
    subjects_list = [(subj.id, f"{subj.code} - {subj.name}") for subj in subjects]

    context = {
        "subjects": subjects_list,   # ✅ replaced SUBJECT_CHOICES
        "dates": valid_dates,
        "schedules": schedules,
    }
    return render(request, "admin_attendance_selector.html", context)


# -------------------------
# Attendance Show & Faculty Assignment
# -------------------------
@login_required(login_url='/admin-login/')
def attendance_show(request):
    if request.user.role != "admin":
        messages.error(request, "Access Denied! Only Admin can access this page.")
        return redirect("student_dashboard")

    # Handle faculty assignment
    if request.method == "POST":
        action = request.POST.get("action")
        if action == "assign":
            sheet_id = request.POST.get("sheet_id")
            faculty_id = request.POST.get("faculty_id")
            sheet = get_object_or_404(AttendanceSheet, id=sheet_id)
            faculty = get_object_or_404(CustomUser, id=faculty_id, role='faculty')
            sheet.assigned_faculty = faculty
            sheet.assigned_at = timezone.now()
            sheet.save()
            messages.success(request, f"Assigned to {faculty.get_full_name() or faculty.username}")
            return redirect(f"{reverse('attendance_show')}?subject={sheet.subject.id}&date={sheet.date}")

    # Get query params
    subject_id = request.GET.get("subject")
    date_param = request.GET.get("date")

    bookings = []
    sheet = None
    date_obj = None
    subject_obj = None  #  store Subject object

    if date_param:
        try:
            date_obj = datetime.strptime(date_param, "%Y-%m-%d").date()
        except ValueError:
            messages.error(request, "Invalid date format. Please use YYYY-MM-DD.")
            date_obj = None

    if subject_id and date_obj:
        #  Get actual Subject object
        subject_obj = get_object_or_404(Subject, id=subject_id)

        # Fetch bookings for this subject and date
        bookings = ExamBooking.objects.filter(
            subject=subject_obj,
            date=date_obj
        ).select_related('student').order_by('student__student_id')

        # Find schedule containing this date
        schedule = ExamDuration.objects.filter(
            start_date__lte=date_obj,
            end_date__gte=date_obj
        ).first()

        if not schedule:
            messages.error(request, "No exam schedule found for the selected date.")
            bookings = []
            sheet = None
        else:
            # Get or create attendance sheet with schedule
            sheet, created = AttendanceSheet.objects.get_or_create(
                subject=subject_obj,
                date=date_obj,
                schedule=schedule
            )
            # Link bookings to attendance sheet
            sheet.bookings.set(bookings)
            sheet.save()

    # Fetch all faculties for dropdown
    faculties = CustomUser.objects.filter(role='faculty')

    context = {
        "subject": subject_obj,   # ✅ now Subject object, not raw string
        "date": date_obj,
        "bookings": bookings,
        "sheet": sheet,
        "faculties": faculties,
    }

    return render(request, "admin_attendance_sheet.html", context)

@login_required(login_url='/admin-login/')
def attendance_export_word(request, sheet_id):
    if request.user.role != "admin":
        messages.error(request, "Access Denied! Only Admin can access this page.")
        return redirect("student_dashboard")

    sheet = get_object_or_404(AttendanceSheet, id=sheet_id)
    bookings = sheet.bookings.select_related("student").order_by("student__student_id")

    # Create Word Document
    doc = Document()

    # ---------- Page Setup ----------
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ---------- Function to Add Page Header ----------
    def add_page_header(doc, page_num):
        # Logo
        logo_path = os.path.join(settings.BASE_DIR, "static", "images", "logo.png")
        if os.path.exists(logo_path):
            header_paragraph = doc.add_paragraph()
            header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            header_paragraph.add_run().add_picture(logo_path, width=Inches(2.5))

        doc.add_paragraph("")  # spacing

        # Subject Info (without faculty)
        info = doc.add_paragraph()
        info.add_run(f"Subject Code: {sheet.subject.code}\n").bold = True
        info.add_run(f"Subject Name: {sheet.subject.name}\n").bold = True
        info.add_run(f"Date: {sheet.date.strftime('%d-%m-%Y')}\n").bold = True
        info.add_run(f"Page: {page_num}").bold = True
        doc.add_paragraph("")

    # ---------- Function to Add Table ----------
    def add_attendance_table(doc, booking_slice, start_index):
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Sr. No"
        hdr_cells[1].text = "Roll No"
        hdr_cells[2].text = "Student Name"
        hdr_cells[3].text = "Signature"

        for idx, b in enumerate(booking_slice, start=start_index):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = b.student.student_id or "-"
            row_cells[2].text = b.student.get_full_name() or b.student.username
            row_cells[3].text = ""

    # ---------- Split Students into Pages ----------
    page_size = 30
    total_students = len(bookings)
    total_pages = (total_students + page_size - 1) // page_size

    for page_num in range(total_pages):
        start = page_num * page_size
        end = start + page_size
        student_slice = bookings[start:end]

        add_page_header(doc, page_num + 1)
        add_attendance_table(doc, student_slice, start + 1)

        # Footer for each page
        doc.add_paragraph("")
        doc.add_paragraph("Total Present: _______")
        doc.add_paragraph("Total Absent: _______")
        doc.add_paragraph("")
        doc.add_paragraph("Faculty Signature: __________________")

        # Add page break except for last page
        if page_num < total_pages - 1:
            doc.add_page_break()

    # ---------- Save to Response ----------
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"attendance_{sheet.subject.code}_{sheet.date}.docx"
    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response
# -----------------------
# Manage subjects
# -----------------------

def manage_subjects(request):
    subjects = Subject.objects.all().order_by('semester', 'code')
    
    if request.method == 'POST':
        form = SubjectForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "Subject added successfully!")
            return redirect('manage_subjects')
    else:
        form = SubjectForm()

    return render(request, 'manage_subjects.html', {'form': form, 'subjects': subjects})

def edit_subject(request, pk):
    subject = get_object_or_404(Subject, pk=pk)
    if request.method == 'POST':
        form = SubjectForm(request.POST, instance=subject)
        if form.is_valid():
            form.save()
            messages.success(request, "Subject updated successfully!")
            return redirect('manage_subjects')
    else:
        form = SubjectForm(instance=subject)

    return render(request, 'edit_subject.html', {'form': form})

def delete_subject(request, pk):
    subject = get_object_or_404(Subject, pk=pk)
    subject.delete()
    messages.success(request, "Subject deleted successfully!")
    return redirect('manage_subjects')

# -----------------------
# Manage Faculty
# -----------------------
@login_required(login_url="/admin-login/")
def manage_faculty(request):
    if request.user.role != "admin":
        messages.error(request, "Access Denied! Only Admin can access this page.")
        return redirect("home")

    pending_faculty = CustomUser.objects.filter(role="faculty", is_approved=False)
    approved_faculty = CustomUser.objects.filter(role="faculty", is_approved=True)

    if request.method == "POST":
        action = request.POST.get("action")
        faculty_id = request.POST.get("faculty_id")
        faculty = get_object_or_404(CustomUser, id=faculty_id)

        if action == "approve":
            faculty.is_approved = True
            faculty.save()
            messages.success(request, f"{faculty.username} has been approved.")

        elif action == "reject":
            faculty.delete()
            messages.warning(request, f"{faculty.username}'s request has been rejected and deleted.")

        elif action == "revoke":
            faculty.is_approved = False
            faculty.save()
            messages.info(request, f"{faculty.username}'s approval has been revoked and moved to pending list.")

        elif action == "delete":
            faculty.delete()
            messages.success(request, f"{faculty.username} has been deleted permanently.")

        return redirect("manage_faculty")

    return render(request, "manage_faculty.html", {
        "pending_faculty": pending_faculty,
        "approved_faculty": approved_faculty,
    })



# -----------------------
# Faculty Dashboard
# -----------------------

@login_required(login_url="/faculty/login/")
def faculty_dashboard(request):
    if request.user.role != "faculty":
        messages.error(request, "Access Denied! Only Faculty can access this page.")
        return redirect("home")

    # Get all AttendanceSheets assigned to this faculty
    assigned_sheets = AttendanceSheet.objects.filter(assigned_faculty=request.user).select_related('subject', 'schedule')

    context = {
        "assigned_sheets": assigned_sheets,
    }

    return render(request, "faculty_dashboard.html", context)



@login_required
def faculty_logout(request):
    logout(request)
    return redirect("faculty_login")


@login_required(login_url='/faculty-login/')
def subject_detail(request, sheet_id):
    sheet = get_object_or_404(AttendanceSheet, id=sheet_id)
    subject = sheet.subject

    # ✅ Get all student bookings for this subject/date
    assignments = sheet.bookings.select_related("student").order_by("student__student_id")

    # ✅ Get student IDs that already have results for this subject/date
    results_student_ids = AdminResult.objects.filter(
        subject_code=subject.code,
        r_date=sheet.date
    ).values_list('student_id', flat=True)

    # ✅ Exclude students whose results already exist
    assignments = [a for a in assignments if a.student.student_id not in results_student_ids]

    context = {
        "sheet": sheet,
        "subject": subject,
        "assignments": assignments,
    }
    return render(request, "subject_detail.html", context)

# -----------------------
# Manage result
# -----------------------



@login_required(login_url='/faculty-login/')
def add_result(request):
    if request.method == "POST":

        sheet_id = request.POST.get("sheet_id")
        subject_code = request.POST.get("subject_code")
        subject_name = request.POST.get("subject_name")
        faculty_name = request.POST.get("faculty_name")

        # Get multiple rows
        student_ids = request.POST.getlist("student_id[]")
        student_names = request.POST.getlist("student_name[]")
        r_dates = request.POST.getlist("r_date[]")
        marks_list = request.POST.getlist("marks[]")
        faculty_name = request.user.username
        # Save each row
        for i in range(len(student_ids)):
            AdminResult.objects.create(
                student_id=student_ids[i],
                student_name=student_names[i],
                r_date=r_dates[i],
                subject_code=subject_code,
                subject_name=subject_name,
                faculty_name=faculty_name,
                out_of_marks=marks_list[i]
            )

        return redirect('subject_detail', sheet_id=sheet_id)   
@login_required(login_url='/admin-login/')
def manage_result(request):
    results = AdminResult.objects.all().order_by("-r_date")
    return render(request, "manage_result.html", {"results": results})

@login_required(login_url='/admin-login/')
def edit_result(request, pk):
    result = get_object_or_404(AdminResult, pk=pk)
    if request.method == "POST":
        form = AdminResultForm(request.POST, instance=result)
        if form.is_valid():
            form.save()
            messages.success(request, "Result updated successfully.")
            return redirect("manage_result")
        else:
            messages.error(request, "Error updating result.")
    else:
        form = AdminResultForm(instance=result)

    return render(request, "edit_result.html", {"form": form, "result": result})


def delete_result(request, pk):
    result = get_object_or_404(AdminResult, pk=pk)
    if request.method == "POST":
        result.delete()
        messages.success(request, "Result deleted successfully.")
        return redirect("manage_result")
    
@login_required(login_url='/admin-login/')
def create_schedule(request):
    if request.user.role != "admin":
        messages.error(request, "Access Denied! Only Admin can access this page.")
        return redirect("student_dashboard")

    # Subjects that have bookings
    subjects = Subject.objects.filter(
        id__in=ExamBooking.objects.values_list('subject', flat=True).distinct()
    )

    selected_subject = None
    selected_date = request.GET.get("date")
    students = []
    schedule = None
    date_obj = None
    total_seats = 0

    # Selected subject
    if request.GET.get("subject"):
        selected_subject = get_object_or_404(Subject, id=request.GET.get("subject"))

    # If both subject and date are selected
    if selected_subject and selected_date:

        # Convert date
        try:
            date_obj = datetime.strptime(selected_date, "%Y-%m-%d").date()
        except ValueError:
            messages.error(request, "Invalid date format.")
            date_obj = None

        if date_obj:

            # Get booked students
            students = ExamBooking.objects.filter(
                subject=selected_subject,
                date=date_obj
            ).select_related("student")

            total_seats = students.count()

            # Create or get exam schedule
            schedule, created = ExamSchedule.objects.get_or_create(
                subject=selected_subject,
                date=date_obj,
                defaults={"start_time": None, "end_time": None}
            )

            # ------------------------------------
            #            POST REQUEST
            # ------------------------------------
            if request.method == "POST":

                # Check if rooms already assigned
                already_assigned = RoomAssignment.objects.filter(
                    subject=selected_subject,
                    date=date_obj
                )
                if already_assigned.exists():
                    messages.error(request, "Rooms are already assigned for this subject and date!")
                    return redirect(f"/create-schedule/?subject={selected_subject.id}&date={selected_date}")




                # Save exam time
                start_time = request.POST.get("start_time")
                end_time = request.POST.get("end_time")

                if start_time:
                    schedule.start_time = start_time
                if end_time:
                    schedule.end_time = end_time
                schedule.save()

                # Assign rooms
                room_number = request.POST.get("room")
                if room_number:
                    student_ids = list(students.values_list("student__id", flat=True))
                    max_per_room = 30
                    chunks = [student_ids[i:i + max_per_room] for i in range(0, len(student_ids), max_per_room)]

                    for idx, chunk in enumerate(chunks):
                        # get selected faculty ID
                        faculty_id = request.POST.get("faculty")

                        try:
                            faculty_obj = CustomUser.objects.get(id=faculty_id, role='faculty')
                        except CustomUser.DoesNotExist:
                            messages.error(request, "Invalid faculty selected!")
                            return redirect(request.path)
                        RoomAssignment.objects.create(
                            subject=selected_subject,
                            date=date_obj,
                            room=f"{room_number}-{idx+1}" if len(chunks) > 1 else room_number,
                            start_time=schedule.start_time,
                            end_time=schedule.end_time,
                            faculty=faculty_obj
                        ).students.set(chunk)

                    messages.success(
                        request,
                        f"{len(student_ids)} students assigned into {len(chunks)} room(s) successfully!"
                    )

                return redirect(f"/create-schedule/?subject={selected_subject.id}&date={selected_date}")


    context = {
        "subjects": subjects,
        "selected_subject": selected_subject,
        "selected_date": selected_date,
        "students": students,
        "schedule": schedule,
        "total_seats": total_seats,
        "faculties": CustomUser.objects.filter(role="faculty"),
    }

    return render(request, "create_schedule.html", context)

def assigned_rooms(request):
    subject_id = request.GET.get("subject")
    date_str = request.GET.get("date")

    subject = None
    date_obj = None

    rooms = RoomAssignment.objects.all().select_related("faculty")

    # Filter by subject
    if subject_id:
        try:
            subject = Subject.objects.get(id=subject_id)
            rooms = rooms.filter(subject=subject)
        except Subject.DoesNotExist:
            subject = None

    # Filter by date
    if date_str:
        try:
            date_obj = datetime.strptime(date_str, "%Y-%m-%d").date()
            rooms = rooms.filter(date=date_obj)
        except ValueError:
            date_obj = None

    context = {
        "rooms": rooms,
        "subject": subject,   # subject object
        "date": date_obj,     # parsed date
        "date_raw": date_str, # original GET date
    }

    return render(request, "assigned_rooms.html", context)


# ---------------------------
# Update Room Assignment
# ---------------------------
def update_schedule(request, schedule_id):
    room = get_object_or_404(RoomAssignment, id=schedule_id)
    
    if request.user.role != "admin":
        messages.error(request, "Access Denied! Only Admin can update this.")
        return redirect('assigned_rooms')

    if request.method == "POST":
        room_number = request.POST.get("room")
        start_time = request.POST.get("start_time")
        end_time = request.POST.get("end_time")

        if room_number:
            room.room = room_number
        if start_time:
            room.start_time = start_time
        if end_time:
            room.end_time = end_time

        room.save()
        messages.success(request, "Room schedule updated successfully!")
        return redirect(f"/assigned-rooms/?subject={room.subject.id}&date={room.date.strftime('%Y-%m-%d')}")

    context = {
        "room": room
    }
    return render(request, "update_schedule.html", context)


# ---------------------------
# Delete Room Assignment
# ---------------------------
def delete_schedule_room(request, schedule_id):
    room = get_object_or_404(RoomAssignment, id=schedule_id)
    
    if request.user.role != "admin":
        messages.error(request, "Access Denied! Only Admin can delete this.")
        return redirect('assigned_rooms')

    subject_id = room.subject.id if room.subject else None
    date_str = room.date.strftime('%Y-%m-%d') if room.date else None

    room.delete()
    messages.success(request, "Room assignment deleted successfully!")
    
    return redirect('assigned_rooms')


def student_assigned_rooms(request):
    student = request.user

    # Subjects whose result is declared
    subjects_with_result = AdminResult.objects.filter(
        student_id=student.student_id
    ).values_list('subject_code', flat=True)

    # Remove student ONLY from rooms where:
    # same subject AND result is declared AND room date <= result date
    for subject_code in subjects_with_result:

        # get result declared date for subject
        result_date = AdminResult.objects.filter(
            student_id=student.student_id,
            subject_code=subject_code
        ).values_list("r_date", flat=True).first()

        if result_date:
            rooms_to_remove = RoomAssignment.objects.filter(
                students=student,
                subject__code=subject_code,
                date__lte=result_date  # exam is definitely completed
            )

            for room in rooms_to_remove:
                room.students.remove(student)

    # Remaining room assignments
    rooms = RoomAssignment.objects.filter(
        students=student
    ).select_related('subject', 'faculty').order_by('date')

    results = AdminResult.objects.filter(student_id=student.student_id)

    context = {
        "rooms": rooms,
        "results": results,
        "pass_mark": 35,
    }
    return render(request, "student_assigned_rooms.html", context)
